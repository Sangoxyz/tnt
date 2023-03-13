import docx
import re
import os
import docx.enum.style
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import  WD_ALIGN_PARAGRAPH

def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def Taode(text1, text2, Num):
    error = False
    listtext1 = re.split(f'\s+',text1)
    listtext2 = re.split(f'\s+',text2)
    document = Document()
    
    #page setting
    sections = document.sections
    section = sections[0]
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(1)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)

    #content
    text = 'KIỂM TRA TIÊU CHUẨN 11 CHUYÊN NGÀNH VTĐ (ĐIỆN SỐ '+ f'{Num}' + ')'
    heading = document.add_paragraph()
    document.styles['Normal']
    heading.style.font.name = 'Times New Roman'
    heading.add_run(text).bold = True
    heading.style.font.size = Pt(14)
    heading.alignment =  WD_ALIGN_PARAGRAPH.CENTER


    para = document.add_paragraph('Họ và tên:.................................\n')
    paraformat = document.styles['Normal'].paragraph_format
    para.style.font.name = 'Times New Roman'
    para.style.font.size = Pt(14)
    para.add_run('Cấp bậc:...........................................\n')
    para.add_run('Chức vụ:...........................................\n')
    para.add_run('Đơn vị:............................................')
    document.add_paragraph('ĐIỆN MÃ:')
    paragraphs = document.paragraphs
    document.paragraphs[2].alignment = WD_ALIGN_PARAGRAPH.CENTER

    table1 = document.add_table(rows=1, cols=6)
    cell_tb1 = table1.rows[0].cells
    cell_tb1[0].text = 'NGHĨA RÕ'
    cell_tb1[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cellNR = cell_tb1[0].paragraphs[0]
    cellNR.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_tb1[1].text = 'MẬT'
    make_rows_bold(table1.rows[0])
    i= 5
    while(i>=2):
        cell_tb1[i-1].merge(cell_tb1[i])
        i -= 1
    cellMat = cell_tb1[1].paragraphs[0]
    cellMat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_tb1[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if len(listtext1) % 5 == 0:
        index = int(len(listtext1) / 5)
    else: index = 1 + int(len(listtext1) / 5)
    for i in range(0,index):
        row_cells = table1.add_row().cells
        row_cells[0].text = ' '.join(listtext1[5*i:(i+1)*5])
        row_cells[0].text = row_cells[0].text.capitalize()
    for cell in table1.columns[0].cells:
        cell.width = Inches(3.5)
    for i in range(1,6):
        for cell in table1.columns[i].cells:
            cell.width = Cm(1.7)
    table1.style = 'Table Grid'
    document.add_paragraph('\nĐIỆN DỊCH:')
    paragraphs = document.paragraphs
    document.paragraphs[3].alignment = WD_ALIGN_PARAGRAPH.CENTER


    table2 = document.add_table(rows=1, cols=6)
    cell_tb2 = table2.rows[0].cells
    cell_tb2[0].text = 'MẬT'
    cell_tb2[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cellNR = cell_tb2[0].paragraphs[0]
    cellNR.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_tb2[5].text = 'NGHĨA RÕ'
    make_rows_bold(table2.rows[0])
    i= 4
    while(i>=1):
        cell_tb2[i-1].merge(cell_tb2[i])
        i -= 1
    cellMat2 = cell_tb2[5].paragraphs[0]
    cellMat2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_tb2[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if len(listtext2) % 5 == 0:
        index = int(len(listtext2) / 5)
    else: index = 1 + int(len(listtext2) / 5)
    x = 0
    for i in range(0,index):
        row_cells = table2.add_row().cells
        for j in range(0,4):
            try:
                row_cells[j].text = listtext2[x]
                x += 1
            except:
                break 
    for cell in table2.columns[5].cells:
        cell.width = Inches(3.5)
    for i in range(0,5):
        for cell in table2.columns[i].cells:
            cell.width = Cm(1.7)
    table2.style = 'Table Grid'
    para2 = document.add_paragraph()
    para2.add_run('\n* Kết quả kiểm tra:').bold =True
    para2.add_run('\n- Thời gian: ..........Phút ......giây;')
    para2.add_run('\n- Sai: ......nhóm; Thiếu: ......nhóm; Thừa: ......nhóm')
    para2.add_run('\n- Điểm: ......; Xếp loại: ........')
    para3 = document.add_paragraph('\nGIÁM KHẢO KIỂM TRA')
    para3.paragraph_format.left_indent  = Cm(9.0)
    para3.add_run('\n \n \n \n')
    para3.add_run('...............................')
    paragraphs = document.paragraphs
    document.paragraphs[5].alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        document.save('De kiem tra tieu chuan 11 (De so '+ f'{Num}' +').docx')
    except:
        error = True
    #os.system('alo.docx')
    return error    
def Laydulieu(text):
    alltext = text.split()
    numcoded = ''
    words = ''
    for NW in alltext:
        if re.match(r'[0-9]{5}', NW):
            numcoded = numcoded  + ''.join(re.findall(r'[0-9]{5}', NW)) + ' '
        else:
            words = words  + ''.join(NW) + ' '
    return(words, numcoded)
        
#Taode('thường dựa vào trục đường chính dải địa hình có giá trị để cơ động lực lượng và triển khai đội hình ; sử dụng mọi biện pháp để che giấu âm moi am am am','53353	32133	50555	03415	47662 73943	36688	38415	97476	94288 15705	81349	95348 69598	49827 94210	25629	03798	16659	73069 21897	81795	15943	69138	34290')
