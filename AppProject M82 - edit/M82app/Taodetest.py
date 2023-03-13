import docx
import os
import docx.enum.style
from tkinter.constants import FALSE
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import  WD_ALIGN_PARAGRAPH
from docx.text.tabstops import TabStop, TabStops

def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

Num = 0

document = Document()

#page setting
sections = document.sections
section = sections[0]
section.left_margin = Cm(3.5)
section.right_margin = Cm(1)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2)

#content
para = document.add_paragraph('Họ và tên:.................................\n')
paraformat = document.styles['Normal'].paragraph_format
para.style.font.name = 'Times New Roman'
run = document.paragraphs
print(run[0].paragraph)
paraformat.line_spacing = Pt(15)
document.add_paragraph('Cấp bậc:...........................................\n')
para.add_run('Chức vụ:...........................................\n')
para.add_run('Đơn vị:............................................')
document.add_paragraph('ĐIỆN MÃ:')
paragraphs = document.paragraphs
document.paragraphs[2].alignment = WD_ALIGN_PARAGRAPH.CENTER

records = ('nácda','caiunawec','ádcnawe')

table = document.add_table(rows=1, cols=6)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'NGHĨA RÕ'
hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
cellNR = hdr_cells[0].paragraphs[0]
cellNR.alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[1].text = 'MẬT'

i= 5
while(i>=2):
    hdr_cells[i-1].merge(hdr_cells[i])
    i -= 1
cellMat = hdr_cells[1].paragraphs[0]
cellMat.alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
for qty in records:
    row_cells = table.add_row().cells
    row_cells[0].text = qty 
for cell in table.columns[0].cells:
    cell.width = Inches(3.5)
for i in range(1,6):
    for cell in table.columns[i].cells:
        cell.width = Cm(1.5)
    


table.style = 'Table Grid'

document.add_page_break()

document.save('alotest.docx')
os.system('alotest.docx')

